#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de génération du Chapitre 2 du rapport technique PFE
WAITLESS-CHU: Conception et Méthodologie

Génère: Analyse fonctionnelle, conception architecture, 
interfaces, et modules techniques.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime


def setup_document_styles(doc):
    """Configure les styles du document."""
    # Style pour les titres de chapitre
    chapter_style = doc.styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
    chapter_font = chapter_style.font
    chapter_font.name = 'Arial'
    chapter_font.size = Pt(16)
    chapter_font.bold = True
    chapter_font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter_style.paragraph_format.space_before = Pt(24)
    chapter_style.paragraph_format.space_after = Pt(18)
    
    # Style pour les titres de section
    section_style = doc.styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Arial'
    section_font.size = Pt(14)
    section_font.bold = True
    section_font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    section_style.paragraph_format.space_before = Pt(18)
    section_style.paragraph_format.space_after = Pt(12)
    
    # Style pour les sous-titres
    subsection_style = doc.styles.add_style('SubsectionTitle', WD_STYLE_TYPE.PARAGRAPH)
    subsection_font = subsection_style.font
    subsection_font.name = 'Arial'
    subsection_font.size = Pt(12)
    subsection_font.bold = True
    subsection_font.color.rgb = RGBColor(0x45, 0xb7, 0xd1)
    subsection_style.paragraph_format.space_before = Pt(12)
    subsection_style.paragraph_format.space_after = Pt(6)
    
    # Style pour le texte normal
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_chapter_intro(doc):
    """Ajoute l'introduction du chapitre."""
    p = doc.add_paragraph()
    p.style = 'ChapterTitle'
    p.add_run("CHAPITRE 2")
    
    p = doc.add_paragraph()
    p.style = 'ChapterTitle'
    p.add_run("CONCEPTION ET MÉTHODOLOGIE")
    
    doc.add_paragraph()
    
    # Introduction du chapitre
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("Introduction")
    
    intro_text = """Ce chapitre détaille la phase de conception du système WAITLESS-CHU, présentant l'analyse fonctionnelle approfondie, l'architecture technique retenue, ainsi que les principes de conception des interfaces utilisateur. La méthodologie de développement et l'organisation du travail sont également explicitées.

La conception constitue une étape cruciale qui transforme les objectifs définis au chapitre précédent en spécifications techniques concrètes. Cette phase permet de valider la faisabilité technique tout en s'assurant que la solution répond parfaitement aux besoins identifiés.

L'approche adoptée privilégie une conception modulaire et extensible, garantissant la maintenabilité et l'évolutivité du système. Les choix architecturaux sont guidés par les exigences de performance, sécurité et expérience utilisateur définies dans la problématique."""
    
    p = doc.add_paragraph(intro_text)


def generate_functional_analysis(doc):
    """Génère la section analyse fonctionnelle."""
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("2.1 Analyse fonctionnelle détaillée")
    
    # 2.1.1 Analyse des besoins fonctionnels
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.1.1 Spécification des besoins par acteur")
    
    functional_text = """L'analyse des besoins fonctionnels révèle trois profils d'utilisateurs distincts avec des exigences spécifiques :

**Besoins Patient :**
• **Simplicité d'utilisation** : Interface intuitive accessible sans formation préalable
• **Accessibilité universelle** : Fonctionnement sur tout dispositif avec navigateur web
• **Transparence totale** : Visibilité en temps réel de la position et temps d'attente estimé
• **Flexibilité temporelle** : Possibilité de quitter et revenir sans perdre sa place
• **Notifications intelligentes** : Alertes au moment optimal pour se présenter
• **Multilinguisme** : Support français, arabe, anglais selon les préférences

**Besoins Personnel Soignant :**
• **Gestion efficace des files** : Interface dédiée pour appeler et gérer les patients
• **Priorisation flexible** : Capacité de gérer les urgences et cas prioritaires
• **Informations patient** : Accès aux informations essentielles (nom, téléphone, priorité)
• **Statistiques temps réel** : Métriques d'activité et performance du service
• **Historique des consultations** : Traçabilité complète des actions effectuées
• **Alertes intelligentes** : Notifications en cas de retard ou situation critique

**Besoins Administrateur :**
• **Vue d'ensemble globale** : Dashboard consolidé de tous les services
• **Gestion du personnel** : CRUD complet des utilisateurs avec attribution des rôles
• **Configuration des services** : Paramétrage des services, horaires, capacités
• **Analytics avancées** : Rapports détaillés et métriques de performance
• **Système d'alertes** : Monitoring proactif et notifications des dysfonctionnements
• **Audit et conformité** : Journalisation complète pour traçabilité réglementaire"""
    
    doc.add_paragraph(functional_text)
    
    # Tableau des besoins fonctionnels
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Acteur', 'Fonction', 'Description', 'Priorité']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    
    # Données fonctionnelles
    functional_data = [
        ("Patient", "Scanner QR", "Rejoindre file via code QR", "Haute"),
        ("Patient", "Consulter position", "Voir position temps réel", "Haute"),
        ("Patient", "Recevoir notifications", "Alertes SMS/Push", "Moyenne"),
        ("Staff", "Appeler patient", "Gestion de la file", "Haute"),
        ("Staff", "Gérer priorités", "Urgences et cas spéciaux", "Haute"),
        ("Staff", "Consulter historique", "Traçabilité actions", "Moyenne"),
        ("Admin", "Gérer personnel", "CRUD utilisateurs", "Haute"),
        ("Admin", "Superviser services", "Vue d'ensemble", "Haute"),
        ("Admin", "Générer rapports", "Analytics et métriques", "Moyenne"),
    ]
    
    for actor, function, description, priority in functional_data:
        row_cells = table.add_row().cells
        row_cells[0].text = actor
        row_cells[1].text = function
        row_cells[2].text = description
        row_cells[3].text = priority
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 2.1.2 Cas d'usage UML
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.1.2 Diagrammes de cas d'usage")
    
    usecase_text = """**Cas d'usage Patient - "Rejoindre une file d'attente" :**

*Acteur principal :* Patient
*Pré-conditions :* Service hospitalier ouvert, QR code affiché
*Scénario nominal :*
1. Le patient scanne le QR code du service
2. Le système affiche le formulaire de saisie
3. Le patient saisit nom et téléphone
4. Le système valide les informations
5. Le système génère un ticket numérique unique
6. Le système calcule la position dans la file
7. Le système affiche la position et temps d'attente estimé
8. Le patient reçoit les notifications de mise à jour

*Scénarios alternatifs :*
- QR code invalide : Message d'erreur et proposition de saisie manuelle
- Service fermé : Information horaires et proposition de notification
- Informations incomplètes : Validation et demande de correction

**Cas d'usage Staff - "Gérer la file d'attente" :**

*Acteur principal :* Personnel soignant
*Pré-conditions :* Authentification réussie, service assigné
*Scénario nominal :*
1. Le staff se connecte à l'interface dédiée
2. Le système affiche la file d'attente du service
3. Le staff consulte les patients en attente
4. Le staff appelle le patient suivant
5. Le système met à jour les positions
6. Le staff marque la fin de consultation
7. Le système synchronise avec tous les clients

**Cas d'usage Admin - "Superviser le système" :**

*Acteur principal :* Administrateur
*Pré-conditions :* Droits administrateur
*Scénario nominal :*
1. L'admin accède au dashboard global
2. Le système affiche l'état de tous les services
3. L'admin consulte les métriques en temps réel
4. L'admin gère les alertes et incidents
5. L'admin génère les rapports d'activité"""
    
    doc.add_paragraph(usecase_text)
    
    # 2.1.3 Exigences non fonctionnelles
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.1.3 Exigences non fonctionnelles")
    
    # Tableau des exigences non fonctionnelles
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Catégorie', 'Exigence', 'Métrique']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    
    # Données exigences
    nfr_data = [
        ("Performance", "Temps de réponse API", "< 200ms en moyenne"),
        ("Performance", "Débit transactions", "100+ req/sec"),
        ("Performance", "Utilisateurs simultanés", "1500+ connexions"),
        ("Disponibilité", "Uptime système", "99.7% (8h/an d'arrêt max)"),
        ("Disponibilité", "Récupération après panne", "< 5 minutes"),
        ("Sécurité", "Authentification", "JWT avec expiration"),
        ("Sécurité", "Chiffrement données", "HTTPS/TLS 1.3"),
        ("Sécurité", "Protection RGPD", "Conformité complète"),
        ("Utilisabilité", "Apprentissage", "< 2 minutes pour nouveau patient"),
        ("Utilisabilité", "Accessibilité", "WCAG 2.1 niveau AA"),
        ("Compatibilité", "Navigateurs", "Chrome, Firefox, Safari, Edge"),
        ("Compatibilité", "Dispositifs", "Desktop, tablette, mobile"),
        ("Scalabilité", "Extension services", "Ajout sans modification code"),
        ("Maintenabilité", "Documentation code", "100% fonctions documentées"),
    ]
    
    for category, requirement, metric in nfr_data:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = requirement
        row_cells[2].text = metric
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)


def generate_architecture_design(doc):
    """Génère la section conception de l'architecture."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("2.2 Conception de l'architecture système")
    
    # 2.2.1 Architecture technique détaillée
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.2.1 Modèle architectural en couches")
    
    architecture_text = """L'architecture du système WAITLESS-CHU repose sur un modèle en couches qui sépare clairement les responsabilités :

**Couche Présentation (Frontend) :**
Cette couche gère l'interface utilisateur et les interactions client. Elle est constituée de :
• **Pages HTML5 sémantiques** : Structure claire avec support accessibility
• **Styles CSS3 modernes** : Design responsive avec variables CSS et animations
• **Scripts JavaScript ES6+** : Logique cliente asynchrone avec gestion d'état
• **APIs Web natives** : Camera API pour QR scanner, LocalStorage pour cache
• **WebSocket Client** : Communication bidirectionnelle temps réel

**Couche Logique Métier (Backend API) :**
Cette couche contient la logique applicative et les règles métier :
• **Routeurs FastAPI** : Endpoints REST organisés par domaine fonctionnel
• **Services métier** : Logique de gestion des files, calculs de priorités
• **Gestionnaire WebSocket** : Synchronisation temps réel multi-clients
• **Générateur QR** : Création de codes QR sécurisés avec validation
• **Système d'authentification** : JWT avec gestion granulaire des rôles

**Couche Accès aux Données (Database) :**
Cette couche assure la persistance et l'intégrité des données :
• **Modèles SQLAlchemy** : ORM avec relations et contraintes d'intégrité
• **Migrations Alembic** : Versioning et évolution du schéma de données
• **Optimisations PostgreSQL** : Index, requêtes optimisées, pooling connexions
• **Logging centralisé** : Audit trail et traçabilité complète des opérations

**Communication Inter-Couches :**
• **REST APIs** : Communication synchrone pour opérations CRUD
• **WebSocket** : Communication asynchrone pour mises à jour temps réel
• **ORM Queries** : Abstraction de l'accès aux données avec type safety
• **Exception Handling** : Gestion d'erreurs propagée à travers les couches"""
    
    doc.add_paragraph(architecture_text)
    
    # 2.2.2 Modélisation des données
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.2.2 Modèle conceptuel de données")
    
    data_model_text = """Le modèle de données WAITLESS-CHU est conçu pour optimiser les performances tout en maintenant l'intégrité référentielle :

**Entité User (Utilisateur) :**
• id (UUID) : Identifiant unique
• email (String, unique) : Adresse email pour authentification
• hashed_password (String) : Mot de passe hashé avec bcrypt
• full_name (String) : Nom complet de l'utilisateur
• phone (String, nullable) : Numéro de téléphone
• role (Enum) : Rôle (PATIENT, STAFF, DOCTOR, ADMIN)
• assigned_service_id (FK, nullable) : Service assigné pour le personnel
• is_active (Boolean) : Statut actif/inactif
• created_at (DateTime) : Date de création

**Entité Service (Service Hospitalier) :**
• id (Integer) : Identifiant numérique
• name (String) : Nom du service (ex: "Cardiologie")
• description (Text, nullable) : Description détaillée
• location (String, nullable) : Localisation physique
• status (Enum) : Statut (ACTIVE, INACTIVE, EMERGENCY)
• priority (Enum) : Priorité par défaut (LOW, MEDIUM, HIGH)
• max_wait_time (Integer) : Temps d'attente maximum autorisé
• avg_wait_time (Float) : Temps d'attente moyen calculé
• created_at (DateTime) : Date de création

**Entité Ticket (Ticket de File) :**
• id (UUID) : Identifiant unique
• ticket_number (String, unique) : Numéro de ticket affiché
• patient_name (String) : Nom du patient
• patient_phone (String) : Téléphone du patient
• service_id (FK) : Référence vers le service
• status (Enum) : Statut (WAITING, COMPLETED, CANCELLED, EXPIRED)
• priority (Enum) : Priorité du ticket
• position_in_queue (Integer) : Position actuelle dans la file
• estimated_wait_time (Integer) : Temps d'attente estimé en minutes
• qr_code (Text) : Code QR base64 encodé
• created_at (DateTime) : Date de création du ticket
• called_at (DateTime, nullable) : Date d'appel du patient
• completed_at (DateTime, nullable) : Date de fin de consultation

**Entité QueueLog (Journal des Files) :**
• id (Integer) : Identifiant auto-incrémenté
• ticket_id (FK, nullable) : Référence vers le ticket concerné
• service_id (FK, nullable) : Référence vers le service
• user_id (FK, nullable) : Utilisateur ayant effectué l'action
• action (String) : Type d'action (join, call, complete, cancel)
• details (Text, nullable) : Détails additionnels de l'action
• timestamp (DateTime) : Horodatage précis de l'action

**Relations et Contraintes :**
• User.assigned_service_id → Service.id (Many-to-One)
• Ticket.service_id → Service.id (Many-to-One)
• QueueLog.ticket_id → Ticket.id (Many-to-One)
• QueueLog.service_id → Service.id (Many-to-One)
• QueueLog.user_id → User.id (Many-to-One)

**Index d'optimisation :**
• Index composite (service_id, status, created_at) sur Ticket
• Index (email) unique sur User
• Index (ticket_number) unique sur Ticket
• Index (timestamp) sur QueueLog pour requêtes temporelles"""
    
    doc.add_paragraph(data_model_text)
    
    # 2.2.3 Sécurité et authentification
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.2.3 Architecture de sécurité")
    
    security_text = """La sécurité du système repose sur une approche multicouche :

**Authentification JWT :**
• **Tokens signés** : Utilisation de clés secrètes pour signature et validation
• **Expiration configurable** : Durée de vie limitée (30 minutes par défaut)
• **Refresh tokens** : Renouvellement automatique pour sessions longues
• **Claims personnalisés** : Inclusion du rôle et permissions dans le token

**Autorisation basée sur les rôles :**
• **RBAC (Role-Based Access Control)** : Contrôle d'accès granulaire
• **Middleware de validation** : Vérification automatique des permissions
• **Isolation des données** : Accès restreint selon le rôle et service assigné
• **Audit trail** : Journalisation de tous les accès et modifications

**Protection des données :**
• **Chiffrement en transit** : HTTPS/TLS 1.3 obligatoire pour toutes les communications
• **Hachage des mots de passe** : bcrypt avec salt automatique et coût adaptatif
• **Validation des entrées** : Schemas Pydantic pour validation stricte côté serveur
• **Sanitisation des sorties** : Prévention des attaques XSS et injection

**Sécurité applicative :**
• **CORS configuré** : Limitation des origines autorisées
• **Rate limiting** : Protection contre les attaques par déni de service
• **Validation des QR codes** : Vérification d'intégrité et horodatage
• **Session management** : Gestion sécurisée des sessions utilisateur

**Conformité réglementaire :**
• **RGPD compliance** : Gestion des données personnelles selon la réglementation
• **Minimisation des données** : Collection uniquement des informations nécessaires
• **Droit à l'oubli** : Procédures de suppression des données sur demande
• **Consentement explicite** : Accord utilisateur pour traitement des données"""
    
    doc.add_paragraph(security_text)


def generate_interface_design(doc):
    """Génère la section conception des interfaces."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("2.3 Conception des interfaces utilisateur")
    
    # 2.3.1 Principes de design
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.3.1 Philosophie et principes de design")
    
    design_principles_text = """La conception des interfaces WAITLESS-CHU s'appuie sur des principes éprouvés d'ergonomie et d'expérience utilisateur :

**Approche Mobile-First :**
• **Design responsive natif** : Conception prioritaire pour écrans mobiles puis adaptation desktop
• **Touch-friendly** : Éléments interactifs dimensionnés pour interaction tactile (44px minimum)
• **Navigation simplifiée** : Hiérarchie claire et chemins de navigation courts
• **Performance mobile** : Optimisation des ressources pour connexions limitées

**Accessibilité universelle :**
• **Contraste élevé** : Ratios de contraste conformes WCAG 2.1 niveau AA (4.5:1 minimum)
• **Navigation clavier** : Tous les éléments accessibles au clavier avec indicateurs focus
• **Textes alternatifs** : Descriptions pour tous les éléments visuels
• **Tailles de police** : Texte dimensionnable jusqu'à 200% sans perte de fonctionnalité

**Simplicité cognitive :**
• **Règle du 7±2** : Limitation du nombre d'éléments par écran
• **Conventions établies** : Respect des patterns d'interface familiers
• **Feedback immédiat** : Retour visuel pour toutes les actions utilisateur
• **États clairs** : Indication visuelle des états (chargement, succès, erreur)

**Cohérence visuelle :**
• **Système de design** : Composants réutilisables avec variables CSS
• **Palette de couleurs** : Couleurs cohérentes avec signification sémantique
• **Typographie hiérarchisée** : 6 niveaux de titres avec proportions harmonieuses
• **Espacement systématique** : Grille basée sur un multiple de 8px"""
    
    doc.add_paragraph(design_principles_text)
    
    # 2.3.2 Charte graphique
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.3.2 Charte graphique et identité visuelle")
    
    # Tableau de la charte graphique
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Élément', 'Valeur', 'Usage']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    
    # Données charte graphique
    design_data = [
        ("Couleur Primaire", "#2c5aa0 (Bleu médical)", "Éléments principaux, CTAs"),
        ("Couleur Secondaire", "#45b7d1 (Bleu clair)", "Éléments secondaires, liens"),
        ("Couleur Accent", "#96ceb4 (Vert apaisant)", "Confirmations, succès"),
        ("Couleur Alerte", "#ff6b6b (Rouge)", "Erreurs, urgences"),
        ("Couleur Neutre", "#f8f9fa (Gris clair)", "Arrière-plans, bordures"),
        ("Police Principale", "Poppins", "Textes interface"),
        ("Police Système", "Arial, sans-serif", "Fallback system"),
        ("Taille Base", "16px", "Texte standard"),
        ("Taille H1", "2.5rem (40px)", "Titres principaux"),
        ("Taille H2", "2rem (32px)", "Titres sections"),
        ("Espacement Base", "8px", "Unité de grille"),
        ("Border Radius", "8px", "Coins arrondis"),
        ("Shadow", "0 2px 10px rgba(0,0,0,0.1)", "Élévation cards"),
    ]
    
    for element, value, usage in design_data:
        row_cells = table.add_row().cells
        row_cells[0].text = element
        row_cells[1].text = value
        row_cells[2].text = usage
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 2.3.3 Wireframes et maquettes
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.3.3 Conception des interfaces principales")
    
    wireframes_text = """**Interface Patient - Scan QR et Ticket :**
L'interface patient privilégie la simplicité extrême avec un workflow en 3 étapes :
1. **Écran de scan** : Vue caméra plein écran avec overlay de détection QR
2. **Formulaire patient** : Saisie nom/téléphone avec validation temps réel
3. **Ticket numérique** : Affichage position, temps estimé avec QR code du ticket

Éléments de design spécifiques :
• **Boutons larges** : CTAs de 56px de hauteur minimum pour facilité tactile
• **Texte contrasté** : Blanc sur fond coloré pour lisibilité optimale
• **Indicateurs de progression** : Barre de progression claire en 3 étapes
• **États de chargement** : Spinners et animations fluides pour feedback

**Interface Staff - Gestion des Files :**
L'interface personnel soignant optimise l'efficacité opérationnelle :
• **Liste patients** : Vue tabulaire avec informations essentielles (nom, priorité, temps attente)
• **Actions rapides** : Boutons "Appeler" et "Terminer" accessibles d'un clic
• **Indicateurs visuels** : Code couleur pour priorités et statuts
• **Métriques temps réel** : Compteurs patients en attente, temps moyen

**Interface Admin - Dashboard Global :**
Le tableau de bord administrateur offre une vue d'ensemble complète :
• **Cards métriques** : KPIs principaux avec évolution graphique
• **Liste des services** : État en temps réel de tous les services
• **Graphiques analytics** : Histogrammes et courbes d'activité
• **Centre d'alertes** : Notifications prioritaires avec actions rapides

**Éléments transversaux :**
• **Navigation principale** : Menu hamburger responsive avec sections claires
• **Messages feedback** : Toasts notifications avec auto-dismiss
• **Modale confirmation** : Dialogues pour actions critiques
• **États vides** : Illustrations et textes explicatifs pour états sans données"""
    
    doc.add_paragraph(wireframes_text)


def generate_technical_modules(doc):
    """Génère la section conception des modules techniques."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("2.4 Conception technique des modules")
    
    # 2.4.1 Module de gestion des files
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.4.1 Algorithme de gestion des files d'attente")
    
    queue_algorithm_text = """Le module de gestion des files constitue le cœur métier du système, implémentant des algorithmes intelligents pour optimiser l'expérience patient :

**Algorithme de positionnement :**
L'attribution de position suit une logique de priorité avec équité temporelle :

```
FONCTION calculer_position(service_id, priorité, timestamp_création):
    // Récupération des tickets en attente triés par priorité puis date
    tickets_attente = REQUÊTE(
        SELECT * FROM tickets 
        WHERE service_id = service_id AND status = 'WAITING'
        ORDER BY priority DESC, created_at ASC
    )
    
    position = 1
    POUR chaque ticket DANS tickets_attente:
        SI ticket.priority < priorité:
            position = INDEX(ticket) + 1
            BREAK
        SINON:
            position = INDEX(ticket) + 2
    
    RETOURNER position
```

**Calcul du temps d'attente estimé :**
L'estimation repose sur l'historique et l'état actuel :

```
FONCTION estimer_temps_attente(service_id, position):
    // Temps moyen par consultation (historique 30 jours)
    temps_moyen = MOYENNE(
        SELECT TIMESTAMPDIFF(called_at, completed_at) 
        FROM tickets 
        WHERE service_id = service_id 
        AND completed_at > NOW() - INTERVAL 30 DAY
    )
    
    // Facteur de correction selon l'heure
    facteur_heure = OBTENIR_facteur_affluence(HEURE_ACTUELLE())
    
    // Estimation finale
    temps_estimé = (position - 1) * temps_moyen * facteur_heure
    
    RETOURNER MAX(temps_estimé, 0)
```

**Gestion des priorités :**
Le système supporte 4 niveaux de priorité avec règles métier :
• **EMERGENCY** : Insertion immédiate en tête de file
• **HIGH** : Insertion avant toutes priorités inférieures
• **MEDIUM** : Priorité standard pour cas normaux
• **LOW** : Consultation différable, positionnée en fin

**Optimisations performance :**
• **Cache Redis** : Mise en cache des positions calculées
• **Calculs incrémentaux** : Mise à jour uniquement des positions affectées
• **Batch updates** : Regroupement des mises à jour pour efficacité
• **Index optimisés** : Index composite (service_id, status, priority, created_at)"""
    
    doc.add_paragraph(queue_algorithm_text)
    
    # 2.4.2 Module QR Code
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.4.2 Système de codes QR sécurisés")
    
    qr_module_text = """Le module QR Code assure la génération, validation et sécurisation des codes :

**Génération des QR codes service :**
Chaque service dispose d'un QR code unique contenant :
```json
{
    "type": "service",
    "service_id": 123,
    "service_name": "Cardiologie",
    "timestamp": 1640995200,
    "signature": "sha256_hash"
}
```

**Génération des QR codes ticket :**
Chaque ticket génère un QR code personnel :
```json
{
    "type": "ticket",
    "ticket_id": "uuid-string",
    "ticket_number": "C001",
    "service_id": 123,
    "timestamp": 1640995200,
    "signature": "sha256_hash"
}
```

**Algorithme de validation :**
```
FONCTION valider_qr_code(qr_data):
    données = JSON_DECODE(qr_data)
    
    // Vérification de la signature
    signature_attendue = HMAC_SHA256(
        données.sans_signature, 
        CLÉ_SECRÈTE_QR
    )
    
    SI données.signature != signature_attendue:
        RETOURNER ERREUR("QR code invalide")
    
    // Vérification de la fraîcheur (24h max)
    SI NOW() - données.timestamp > 86400:
        RETOURNER ERREUR("QR code expiré")
    
    // Vérification de l'existence du service
    SI données.type == "service":
        service = CHERCHER_SERVICE(données.service_id)
        SI service.status != "ACTIVE":
            RETOURNER ERREUR("Service fermé")
    
    RETOURNER SUCCÈS(données)
```

**Sécurisation :**
• **Signature HMAC** : Intégrité garantie par signature cryptographique
• **Expiration temporelle** : Validité limitée à 24h pour les QR services
• **Validation côté serveur** : Vérifications multiples avant acceptation
• **Rotation des clés** : Renouvellement périodique des clés de signature"""
    
    doc.add_paragraph(qr_module_text)
    
    # 2.4.3 Module temps réel
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("2.4.3 Architecture WebSocket temps réel")
    
    websocket_text = """Le module temps réel assure la synchronisation instantanée entre tous les clients :

**Gestionnaire de connexions WebSocket :**
```python
class WebSocketManager:
    def __init__(self):
        # Connexions par service pour diffusion ciblée
        self.service_connections: Dict[int, List[WebSocket]] = {}
        # Connexions par ticket pour notifications personnelles
        self.ticket_connections: Dict[str, List[WebSocket]] = {}
        # Connexions admin pour supervision globale
        self.admin_connections: List[WebSocket] = []
    
    async def diffuser_mise_a_jour_file(self, service_id: int, données: dict):
        """Diffuse les mises à jour à tous les clients du service"""
        connexions = self.service_connections.get(service_id, [])
        message = {
            "type": "queue_update",
            "service_id": service_id,
            "data": données,
            "timestamp": datetime.utcnow().isoformat()
        }
        
        # Diffusion parallèle pour performance
        await asyncio.gather(*[
            self.envoyer_message_securise(ws, message)
            for ws in connexions
        ])
```

**Types de messages temps réel :**
• **queue_update** : Mise à jour de l'état de la file d'attente
• **position_change** : Changement de position pour un ticket spécifique
• **patient_called** : Notification d'appel d'un patient
• **service_status** : Changement de statut d'un service (ouvert/fermé/urgence)
• **system_alert** : Alertes système pour les administrateurs

**Gestion de la fiabilité :**
• **Heartbeat/Ping-Pong** : Vérification périodique des connexions actives
• **Reconnexion automatique** : Tentatives exponentielles de reconnexion côté client
• **Buffer de messages** : Conservation temporaire des messages en cas de déconnexion
• **Synchronisation à la reconnexion** : Mise à jour de l'état complet à la reconnexion

**Optimisations performance :**
• **Diffusion groupée** : Regroupement des messages pour efficacité réseau
• **Compression** : Compression des messages volumineux
• **Filtrages ciblés** : Envoi uniquement aux clients concernés
• **Rate limiting** : Protection contre le spam de messages"""
    
    doc.add_paragraph(websocket_text)


def generate_chapter_conclusion(doc):
    """Ajoute la conclusion du chapitre."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("Conclusion du chapitre")
    
    conclusion_text = """Ce chapitre a détaillé la phase de conception du système WAITLESS-CHU, transformant les objectifs fonctionnels en spécifications techniques concrètes et implémentables.

L'analyse fonctionnelle approfondie a permis d'identifier précisément les besoins de chaque type d'utilisateur et de formaliser ces besoins en exigences fonctionnelles et non fonctionnelles mesurables. Les diagrammes de cas d'usage UML structurent les interactions et valident la complétude de la couverture fonctionnelle.

La conception architecturale adopte une approche en couches qui sépare clairement les responsabilités tout en optimisant les performances et la maintenabilité. Le modèle de données relationnel équilibre normalisation et performance, tandis que l'architecture de sécurité multicouche garantit la protection des données sensibles.

La conception des interfaces utilisateur privilégie l'accessibilité universelle et l'expérience utilisateur optimale, avec une approche Mobile-First qui assure un fonctionnement optimal sur tous les dispositifs. La charte graphique cohérente renforce l'identité visuelle et la professionellement du système.

Les modules techniques conçus (gestion des files, QR codes, temps réel) implémentent des algorithmes intelligents qui automatisent efficacement les processus métier tout en conservant la flexibilité nécessaire aux cas d'usage spéciaux.

Cette conception solide constitue la fondation technique qui guidera l'implémentation détaillée présentée dans les chapitres suivants, garantissant que la solution finale répondra parfaitement aux objectifs fixés."""
    
    doc.add_paragraph(conclusion_text)


def generate_chapter2_report():
    """Fonction principale de génération du chapitre 2."""
    print("🚀 Génération du Chapitre 2: Conception et Méthodologie...")
    
    # Créer le document
    doc = Document()
    
    # Configurer les styles
    setup_document_styles(doc)
    
    # Générer chaque section
    print("📋 Génération de l'introduction du chapitre...")
    add_chapter_intro(doc)
    
    print("🔍 Génération de l'analyse fonctionnelle...")
    generate_functional_analysis(doc)
    
    print("🏗️ Génération de la conception architecture...")
    generate_architecture_design(doc)
    
    print("🎨 Génération de la conception interfaces...")
    generate_interface_design(doc)
    
    print("⚙️ Génération des modules techniques...")
    generate_technical_modules(doc)
    
    print("✍️ Génération de la conclusion du chapitre...")
    generate_chapter_conclusion(doc)
    
    # Sauvegarder le document
    filename = "chapitre2.docx"
    doc.save(filename)
    
    print(f"✅ Chapitre 2 généré avec succès: {filename}")
    print(f"📄 Pages générées: ~15 pages")
    print(f"📊 Contenu: Analyse fonctionnelle, architecture, interfaces, modules techniques")
    
    return filename


if __name__ == "__main__":
    generate_chapter2_report()