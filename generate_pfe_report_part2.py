#!/usr/bin/env python3
"""
Professional PFE Report Generator - Part 2
Chapters 1-4 of the technical report
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import os

class ChapterReportGenerator:
    def __init__(self):
        self.document = Document()
        self.setup_document_styles()
        
    def setup_document_styles(self):
        """Setup professional document styles"""
        # Title style
        title_style = self.document.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Times New Roman'
        title_font.size = Pt(18)
        title_font.bold = True
        title_font.color.rgb = RGBColor(0, 0, 0)
        
        # Heading 1 style
        h1_style = self.document.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_font = h1_style.font
        h1_font.name = 'Times New Roman'
        h1_font.size = Pt(16)
        h1_font.bold = True
        h1_font.color.rgb = RGBColor(44, 90, 160)
        
        # Heading 2 style
        h2_style = self.document.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_font = h2_style.font
        h2_font.name = 'Times New Roman'
        h2_font.size = Pt(14)
        h2_font.bold = True
        h2_font.color.rgb = RGBColor(69, 183, 209)
        
        # Heading 3 style
        h3_style = self.document.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_font = h3_style.font
        h3_font.name = 'Times New Roman'
        h3_font.size = Pt(12)
        h3_font.bold = True
        
        # Normal text style
        normal_style = self.document.styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Times New Roman'
        normal_font.size = Pt(11)
        
    def add_title(self, text, level=1):
        """Add a title with appropriate styling"""
        if level == 1:
            paragraph = self.document.add_paragraph(text, style='CustomTitle')
        elif level == 2:
            paragraph = self.document.add_paragraph(text, style='CustomHeading1')
        elif level == 3:
            paragraph = self.document.add_paragraph(text, style='CustomHeading2')
        else:
            paragraph = self.document.add_paragraph(text, style='CustomHeading3')
        
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return paragraph
        
    def add_heading(self, text, level=1):
        """Add a heading with appropriate styling"""
        if level == 1:
            paragraph = self.document.add_paragraph(text, style='CustomHeading1')
        elif level == 2:
            paragraph = self.document.add_paragraph(text, style='CustomHeading2')
        else:
            paragraph = self.document.add_paragraph(text, style='CustomHeading3')
        return paragraph
        
    def add_paragraph(self, text):
        """Add a normal paragraph"""
        return self.document.add_paragraph(text, style='CustomNormal')
        
    def add_image_placeholder(self, caption, width=6):
        """Add a placeholder for an image"""
        paragraph = self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        placeholder = paragraph.add_run(f"[IMAGE PLACEHOLDER: {caption}]")
        placeholder.font.color.rgb = RGBColor(128, 128, 128)
        placeholder.font.italic = True
        
        caption_para = self.document.add_paragraph(f"Figure: {caption}")
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.style = 'CustomNormal'
        
        return paragraph
        
    def add_table(self, headers, rows):
        """Add a professional table"""
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            header_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
            
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
                row_cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                
        return table
        
    def add_list(self, items, numbered=False):
        """Add a bulleted or numbered list"""
        for item in items:
            if numbered:
                paragraph = self.document.add_paragraph(f"{items.index(item) + 1}. {item}", style='CustomNormal')
            else:
                paragraph = self.document.add_paragraph(f"• {item}", style='CustomNormal')
        return paragraph
        
    def add_page_break(self):
        """Add a page break"""
        self.document.add_page_break()
        
    def generate_chapter1(self):
        """Generate Chapter 1: General Project Context"""
        self.add_title("CHAPTER 1: GENERAL PROJECT CONTEXT", 1)
        self.document.add_paragraph()
        
        # Introduction
        self.add_heading("Introduction", 2)
        intro_text = """This first chapter presents the general context in which our WAITLESS-CHU project fits. We will detail the organizational environment, the identified problematics, as well as the objectives and project planning."""
        self.add_paragraph(intro_text)
        self.document.add_paragraph()
        
        # 1.1 Project Environment Presentation
        self.add_heading("1.1 Project Environment Presentation", 2)
        
        # 1.1.1 Institutional Framework
        self.add_heading("1.1.1 Institutional Framework", 3)
        framework_text = """The WAITLESS-CHU project is part of a Final Year Project (PFE) carried out within a higher education institution, in conceptual partnership with Moroccan University Hospital Centers."""
        self.add_paragraph(framework_text)
        self.document.add_paragraph()
        
        # 1.1.2 Health and Digital Context
        self.add_heading("1.1.2 Health and Digital Context", 3)
        context_text = """Moroccan healthcare institutions, particularly CHUs, welcome thousands of patients daily. Traditional queue management presents several challenges:

• Service overload: Long and uncomfortable physical queues
• Lack of visibility: Absence of information on estimated waiting times
• Organizational inefficiency: Manual management prone to errors
• Degraded patient experience: Stress and dissatisfaction related to uncertainty"""
        self.add_paragraph(context_text)
        self.document.add_paragraph()
        
        # 1.1.3 Project Information Sheet
        self.add_heading("1.1.3 Project Information Sheet", 3)
        self.add_paragraph("Table 1.1: Organization Information Sheet")
        
        headers = ["Criterion", "Information"]
        rows = [
            ["Project Name", "WAITLESS-CHU"],
            ["Type", "Intelligent Queue Management System"],
            ["Sector", "Public Health / Hospital Technology"],
            ["Beneficiaries", "Patients, Medical Staff, Administrators"],
            ["Platform", "Web (Multi-device)"],
            ["Development Duration", "6 months"],
            ["Team", "2 student developers"]
        ]
        self.add_table(headers, rows)
        self.document.add_paragraph()
        
        # 1.2 General Project Context
        self.add_heading("1.2 General Project Context", 2)
        
        # 1.2.1 Identified Problematics
        self.add_heading("1.2.1 Identified Problematics", 3)
        problematics_text = """Analysis of the current hospital context reveals several critical issues:

For patients:
• Prolonged waiting times without visibility
• Need to remain physically present
• Stress and uncertainty about deadlines
• Health risks related to gatherings

For medical staff:
• Complex manual queue management
• Difficulties in prioritizing emergencies
• Absence of analytical tools
• Administrative overload

For the institution:
• Operational inefficiency
• Degraded patient satisfaction
• Lack of data for optimization
• Affected institutional image"""
        self.add_paragraph(problematics_text)
        self.document.add_paragraph()
        
        # 1.2.2 Technological Opportunities
        self.add_heading("1.2.2 Technological Opportunities", 3)
        opportunities_text = """Recent technology evolution offers exceptional opportunities:

• Smartphone democratization: 95% equipment rate
• Mature QR technology: Massive post-COVID adoption
• Accessible cloud computing: Scalable infrastructure
• Modern frameworks: Rapid and robust development"""
        self.add_paragraph(opportunities_text)
        self.document.add_paragraph()
        
        # 1.3 WAITLESS-CHU System Architecture
        self.add_heading("1.3 WAITLESS-CHU System Architecture", 2)
        
        # 1.3.1 Architectural Overview
        self.add_heading("1.3.1 Architectural Overview", 3)
        self.add_image_placeholder("General system architecture of WAITLESS-CHU")
        
        arch_text = """The WAITLESS-CHU system adopts a modern three-layer architecture:

Presentation Layer (Frontend):
• Responsive web interface (HTML5/CSS3/JavaScript)
• Multi-device support (desktop, tablet, mobile)
• Integrated QR scanner
• Real-time notifications

Business Logic Layer (Backend):
• RESTful API (FastAPI)
• JWT authentication
• WebSocket management
• Intelligent queue engine

Data Layer (Database):
• PostgreSQL database
• Optimized relational models
• Complete logging
• Automatic backup"""
        self.add_paragraph(arch_text)
        self.document.add_paragraph()
        
        # 1.3.2 Simplified Patient Flow
        self.add_heading("1.3.2 Simplified Patient Flow", 3)
        self.add_image_placeholder("Patient flow diagram")
        
        flow_text = """The patient journey is broken down into simple steps:

1. Arrival → Service QR code scan
2. Registration → Minimal information entry
3. Assignment → Digital ticket generation
4. Tracking → Real-time position notifications
5. Consultation → Automatic call"""
        self.add_paragraph(flow_text)
        self.document.add_paragraph()
        
        # 1.4 Problematic and Proposed Solution
        self.add_heading("1.4 Problematic and Proposed Solution", 2)
        
        # 1.4.1 Problem Statement
        self.add_heading("1.4.1 Problem Statement", 3)
        problem_text = """How to modernize hospital queue management by eliminating physical waiting while providing real-time visibility and advanced management tools for medical staff?"""
        self.add_paragraph(problem_text)
        self.document.add_paragraph()
        
        # 1.4.2 Solution Hypotheses
        self.add_heading("1.4.2 Solution Hypotheses", 3)
        hypotheses_text = """Our approach is based on three fundamental hypotheses:

1. QR Technology: Rapid adoption without application installation
2. Real-time: Significant improvement in patient experience
3. Analytical data: Continuous process optimization"""
        self.add_paragraph(hypotheses_text)
        self.document.add_paragraph()
        
        # 1.4.3 System Objectives
        self.add_heading("1.4.3 System Objectives", 3)
        
        functional_obj = """Functional objectives:
• Reduce perceived waiting time by 70%
• Eliminate 100% of physical waiting
• Provide real-time visibility
• Automate queue management"""
        self.add_paragraph(functional_obj)
        
        technical_obj = """Technical objectives:
• Scalable architecture (1000+ simultaneous users)
• Response time < 200ms
• 99.9% availability
• Enhanced security (GDPR compliant)"""
        self.add_paragraph(technical_obj)
        
        organizational_obj = """Organizational objectives:
• Improve operational efficiency
• Reduce administrative burden
• Provide advanced analytics
• Modernize institutional image"""
        self.add_paragraph(organizational_obj)
        self.document.add_paragraph()
        
        # 1.5 Work to be Done
        self.add_heading("1.5 Work to be Done", 2)
        
        # 1.5.1 Learning Phases
        self.add_heading("1.5.1 Learning Phases", 3)
        learning_text = """The project required acquiring skills in several areas:

Phase 1: Backend Technologies (2 weeks)
• FastAPI and REST API development
• PostgreSQL and data modeling
• JWT authentication and security
• WebSockets and real-time communication

Phase 2: Frontend Technologies (2 weeks)
• Modern HTML5/CSS3 and responsive design
• JavaScript ES6+ and asynchronous programming
• API Fetch and state management
• QR scanner and camera access

Phase 3: Integration and DevOps (1 week)
• Client-server architecture
• Error handling and monitoring
• Automated testing
• Deployment and configuration"""
        self.add_paragraph(learning_text)
        self.document.add_paragraph()
        
        # 1.5.2 Development Phase
        self.add_heading("1.5.2 Development Phase", 3)
        
        sprint1_text = """Sprint 1 - Infrastructure (3 weeks):
• Development environment configuration
• Database modeling
• Authentication API
• Basic administration interface"""
        self.add_paragraph(sprint1_text)
        
        sprint2_text = """Sprint 2 - Core Features (4 weeks):
• Queue management system
• QR code generation and scanning
• Ticket management
• Real-time notifications"""
        self.add_paragraph(sprint2_text)
        
        sprint3_text = """Sprint 3 - User Interfaces (3 weeks):
• Administrator dashboard
• Secretary interface
• Home page and QR scan
• Responsive design"""
        self.add_paragraph(sprint3_text)
        
        sprint4_text = """Sprint 4 - Optimization and Testing (2 weeks):
• Performance testing
• Database optimizations
• Documentation
• Deployment"""
        self.add_paragraph(sprint4_text)
        self.document.add_paragraph()
        
        # 1.5.3 Gantt Chart
        self.add_heading("1.5.3 Gantt Chart", 3)
        self.add_image_placeholder("Project Gantt chart")
        
        # Chapter Conclusion
        self.add_heading("Chapter Conclusion", 2)
        conclusion_text = """This first chapter has established the general context of the WAITLESS-CHU project, highlighting the problematic of hospital queues and the technological opportunity that our solution represents. The proposed architecture and detailed planning constitute solid foundations for the design and development phases that will follow.

The next chapter will address the detailed system design, including the adopted development methodology and the technical environment set up."""
        self.add_paragraph(conclusion_text)
        self.add_page_break()
        
    def save_document(self, filename):
        """Save the document"""
        self.document.save(filename)
        print(f"Document saved as: {filename}")

def main():
    """Main function to generate Chapter 1"""
    generator = ChapterReportGenerator()
    
    print("Generating Chapter 1...")
    generator.generate_chapter1()
    
    # Save the document
    generator.save_document("WAITLESS_CHU_PFE_Report_Chapter1.docx")
    
    print("Chapter 1 completed successfully!")

if __name__ == "__main__":
    main() 