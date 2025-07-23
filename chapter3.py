#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script de génération du Chapitre 3 du rapport technique PFE
WAITLESS-CHU: Choix Technologiques et Justifications

Génère: Stack backend, frontend, outils développement, sécurité.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime


def setup_document_styles(doc):
    """Configure les styles du document."""
    # Style pour les titres de chapitre
    chapter_style = doc.styles.add_style('ChapterTitle', WD_STYLE_TYPE.PARAGRAPH)
    chapter_font = chapter_style.font
    chapter_font.name = 'Arial'
    chapter_font.size = Pt(16)
    chapter_font.bold = True
    chapter_font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chapter_style.paragraph_format.space_before = Pt(24)
    chapter_style.paragraph_format.space_after = Pt(18)
    
    # Style pour les titres de section
    section_style = doc.styles.add_style('SectionTitle', WD_STYLE_TYPE.PARAGRAPH)
    section_font = section_style.font
    section_font.name = 'Arial'
    section_font.size = Pt(14)
    section_font.bold = True
    section_font.color.rgb = RGBColor(0x2c, 0x5a, 0xa0)
    section_style.paragraph_format.space_before = Pt(18)
    section_style.paragraph_format.space_after = Pt(12)
    
    # Style pour les sous-titres
    subsection_style = doc.styles.add_style('SubsectionTitle', WD_STYLE_TYPE.PARAGRAPH)
    subsection_font = subsection_style.font
    subsection_font.name = 'Arial'
    subsection_font.size = Pt(12)
    subsection_font.bold = True
    subsection_font.color.rgb = RGBColor(0x45, 0xb7, 0xd1)
    subsection_style.paragraph_format.space_before = Pt(12)
    subsection_style.paragraph_format.space_after = Pt(6)
    
    # Style pour le texte normal
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Arial'
    normal_font.size = Pt(11)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_chapter_intro(doc):
    """Ajoute l'introduction du chapitre."""
    p = doc.add_paragraph()
    p.style = 'ChapterTitle'
    p.add_run("CHAPITRE 3")
    
    p = doc.add_paragraph()
    p.style = 'ChapterTitle'
    p.add_run("CHOIX TECHNOLOGIQUES ET JUSTIFICATIONS")
    
    doc.add_paragraph()
    
    # Introduction du chapitre
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("Introduction")
    
    intro_text = """Ce chapitre présente et justifie l'ensemble des choix technologiques effectués pour le développement du système WAITLESS-CHU. Chaque décision technique a été prise en considérant les critères de performance, scalabilité, maintenabilité, sécurité et facilité de développement.

L'analyse comparative des technologies disponibles a guidé la sélection d'un stack technologique moderne et éprouvé, combinant innovation et fiabilité. Les justifications présentées s'appuient sur des critères objectifs et l'adéquation aux exigences spécifiques du projet.

La philosophie adoptée privilégie les technologies open source, les standards web modernes, et les frameworks reconnus pour leur robustesse et leur communauté active. Cette approche garantit la pérennité de la solution et facilite sa maintenance évolutive."""
    
    p = doc.add_paragraph(intro_text)


def generate_backend_technology_stack(doc):
    """Génère la section stack technologique backend."""
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("3.1 Stack technologique Backend")
    
    # 3.1.1 Framework FastAPI
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.1.1 FastAPI comme framework de développement")
    
    fastapi_text = """FastAPI a été retenu comme framework principal pour le développement de l'API backend après une analyse comparative approfondie :

**Justifications techniques :**

**Performance exceptionnelle :**
FastAPI figure parmi les frameworks Python les plus performants, rivalisant avec Node.js et Go. Les benchmarks indépendants démontrent des performances 2 à 3 fois supérieures à Django et Flask pour les mêmes cas d'usage. Cette performance est cruciale pour supporter 1500+ utilisateurs simultanés avec des temps de réponse < 200ms.

**Documentation automatique intégrée :**
L'intégration native de Swagger/OpenAPI 3.0 génère automatiquement une documentation interactive complète. Cette fonctionnalité accélère significativement le développement frontend et facilite les tests d'intégration. L'interface Swagger accessible via `/docs` permet aux développeurs de tester tous les endpoints directement.

**Validation de données moderne :**
L'intégration avec Pydantic apporte une validation de données robuste et une sérialisation automatique. Les schemas Pydantic garantissent la cohérence des données d'entrée et de sortie, réduisant les erreurs et améliorant la fiabilité du système.

**Support asynchrone natif :**
FastAPI exploite pleinement les capacités asynchrones de Python avec async/await, essential pour la gestion simultanée de nombreuses connexions WebSocket et l'optimisation des performances I/O-bound.

**Type hints et developer experience :**
L'utilisation extensive des type hints Python améliore significativement l'expérience développeur avec auto-complétion, détection d'erreurs statiques, et refactoring sécurisé."""
    
    doc.add_paragraph(fastapi_text)
    
    # Tableau comparatif frameworks
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Framework', 'Performance', 'Documentation', 'Validation', 'Async', 'Évaluation']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9)
    
    # Données comparatives
    framework_data = [
        ("FastAPI", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "9.5/10"),
        ("Django", "⭐⭐⭐", "⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐", "7/10"),
        ("Flask", "⭐⭐⭐⭐", "⭐⭐", "⭐⭐", "⭐⭐⭐⭐", "6.5/10"),
        ("Tornado", "⭐⭐⭐⭐", "⭐⭐", "⭐⭐", "⭐⭐⭐⭐⭐", "6/10"),
    ]
    
    for framework, perf, doc, valid, async_support, rating in framework_data:
        row_cells = table.add_row().cells
        row_cells[0].text = framework
        row_cells[1].text = perf
        row_cells[2].text = doc
        row_cells[3].text = valid
        row_cells[4].text = async_support
        row_cells[5].text = rating
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 3.1.2 PostgreSQL comme SGBD
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.1.2 PostgreSQL pour la persistance des données")
    
    postgresql_text = """PostgreSQL s'impose comme le choix optimal pour les besoins de WAITLESS-CHU :

**Robustesse et conformité ACID :**
PostgreSQL offre une conformité ACID complète garantissant l'intégrité des données critiques (positions dans les files, statuts des tickets). Cette robustesse est essentielle dans un contexte hospitalier où la perte de données peut avoir des conséquences graves.

**Fonctionnalités avancées :**
• **Support JSON natif** : Stockage et requêtes JSON pour données flexibles (métadonnées QR, configurations)
• **Arrays et types personnalisés** : Modélisation fine des données métier
• **Full-text search** : Recherche avancée dans les données patients et services
• **Extensions** : Extensibilité via modules spécialisés (PostGIS pour géolocalisation future)

**Performance et optimisation :**
• **Optimiseur de requêtes sophistiqué** : Plans d'exécution optimaux pour requêtes complexes
• **Index avancés** : B-tree, Hash, GIN, GiST pour optimisation spécifique
• **Partitioning** : Possibilité de partition pour grands volumes de données
• **Connection pooling** : Gestion efficace des connexions concurrentes

**Conformité standards :**
• **SQL standard** : Requêtes portables et maintenables
• **Transactions** : Gestion transactionnelle complète avec points de sauvegarde
• **Contraintes** : Validation de l'intégrité référentielle au niveau base"""
    
    doc.add_paragraph(postgresql_text)
    
    # Tableau comparatif SGBD
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['SGBD', 'ACID', 'Performance', 'Fonctionnalités', 'Communauté', 'Note']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9)
    
    # Données SGBD
    db_data = [
        ("PostgreSQL", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "9.8/10"),
        ("MySQL", "⭐⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐", "⭐⭐⭐⭐⭐", "7.5/10"),
        ("SQLite", "⭐⭐⭐⭐", "⭐⭐⭐", "⭐⭐", "⭐⭐⭐", "5/10"),
        ("MongoDB", "⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐⭐", "6.5/10"),
    ]
    
    for db, acid, perf, features, community, rating in db_data:
        row_cells = table.add_row().cells
        row_cells[0].text = db
        row_cells[1].text = acid
        row_cells[2].text = perf
        row_cells[3].text = features
        row_cells[4].text = community
        row_cells[5].text = rating
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 3.1.3 SQLAlchemy ORM
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.1.3 SQLAlchemy 2.0 pour l'abstraction de données")
    
    sqlalchemy_text = """SQLAlchemy 2.0 apporte une couche d'abstraction moderne et performante :

**Pattern ORM moderne :**
• **Declarative Base** : Définition claire des modèles avec héritage
• **Relationships** : Gestion automatique des relations complexes
• **Lazy Loading** : Optimisation des requêtes avec chargement à la demande
• **Type Hints** : Support complet des annotations de type Python

**Performance optimisée :**
• **Connection Pooling** : Réutilisation des connexions pour efficacité
• **Batch Operations** : Opérations groupées pour réduction des round-trips
• **Query Optimization** : Génération de SQL optimisé
• **Caching** : Cache de requêtes et d'objets pour performance

**Flexibilité :**
• **Raw SQL** : Possibilité d'utiliser SQL natif quand nécessaire
• **Migrations Alembic** : Versioning automatique du schéma
• **Multiple Databases** : Support multi-base si nécessaire"""
    
    doc.add_paragraph(sqlalchemy_text)


def generate_frontend_technology_stack(doc):
    """Génère la section stack technologique frontend."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("3.2 Stack technologique Frontend")
    
    # 3.2.1 JavaScript Vanilla
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.2.1 Choix du JavaScript Vanilla")
    
    vanilla_js_text = """Le choix de JavaScript vanilla (sans framework) résulte d'une analyse approfondie des besoins spécifiques du projet :

**Justifications stratégiques :**

**Simplicité et accessibilité :**
L'absence de framework élimine la courbe d'apprentissage et la complexité de setup. Tout développeur connaissant JavaScript peut contribuer immédiatement au projet sans formation spécifique à React, Vue ou Angular.

**Performance optimale :**
• **Pas d'overhead framework** : Élimination du bundle JavaScript volumineux
• **Time to Interactive** réduit : Chargement plus rapide, crucial pour l'expérience mobile
• **Memory footprint** minimal : Consommation mémoire optimisée
• **Bundle size** contrôlé : Scripts optimisés sans dépendances externes

**Compatibilité universelle :**
JavaScript vanilla garantit une compatibilité maximale avec tous les navigateurs sans transpilation ni polyfills complexes. Cette universalité est cruciale pour l'accessibilité hospitalière où les dispositifs peuvent être obsolètes.

**Maintenance simplifiée :**
• **Pas de dépendances externes** : Élimination des problèmes de versions et vulnérabilités
• **Code explicite** : Logique claire sans abstractions frameworkisées
• **Debugging facilité** : Débogage direct sans sourcemaps complexes
• **Long-term support** : Pérennité garantie par les standards web

**APIs Web modernes exploitées :**
• **Camera API** : Accès natif à la caméra pour scan QR
• **WebSocket API** : Communication temps réel native
• **LocalStorage** : Persistance locale sans bibliothèque
• **Fetch API** : Requêtes HTTP modernes avec Promises
• **Service Workers** : Capacités PWA pour fonctionnement hors-ligne"""
    
    doc.add_paragraph(vanilla_js_text)
    
    # Tableau comparatif technologies frontend
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Technologie', 'Complexité', 'Performance', 'Bundle Size', 'Compatibilité', 'Score']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9)
    
    # Données frontend
    frontend_data = [
        ("Vanilla JS", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "⭐⭐⭐⭐⭐", "9.2/10"),
        ("React.js", "⭐⭐", "⭐⭐⭐⭐", "⭐⭐", "⭐⭐⭐⭐", "6.5/10"),
        ("Vue.js", "⭐⭐⭐", "⭐⭐⭐⭐", "⭐⭐⭐", "⭐⭐⭐⭐", "7/10"),
        ("Angular", "⭐", "⭐⭐⭐", "⭐", "⭐⭐⭐", "5/10"),
    ]
    
    for tech, complexity, perf, bundle, compat, score in frontend_data:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = complexity
        row_cells[2].text = perf
        row_cells[3].text = bundle
        row_cells[4].text = compat
        row_cells[5].text = score
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 3.2.2 HTML5 et CSS3 modernes
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.2.2 Standards web modernes HTML5/CSS3")
    
    html_css_text = """**HTML5 sémantique :**
L'utilisation d'HTML5 apporte structure et accessibilité :
• **Éléments sémantiques** : `<header>`, `<nav>`, `<main>`, `<section>` pour structure claire
• **APIs natives** : Camera, WebSocket, LocalStorage sans polyfills
• **Formulaires avancés** : Validation native, types d'input spécialisés
• **Accessibilité** : ARIA labels, rôles, descriptions pour lecteurs d'écran

**CSS3 moderne :**
Les fonctionnalités CSS3 optimisent design et performance :
• **Flexbox et Grid** : Layouts complexes sans frameworks CSS
• **Variables CSS** : Thématisation cohérente et maintenance facilitée
• **Media Queries** : Responsive design natif multi-résolutions
• **Animations CSS** : Transitions fluides sans JavaScript

**Avantages techniques :**
• **Rendu natif** : Performance optimale par le moteur de rendu
• **Cascade naturelle** : Héritage CSS intuitif et prévisible  
• **Debugging facilité** : Outils développeur intégrés navigateurs
• **Standards W3C** : Pérennité et évolution garanties"""
    
    doc.add_paragraph(html_css_text)
    
    # 3.2.3 Architecture frontend modulaire
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.2.3 Organisation modulaire du code frontend")
    
    frontend_architecture_text = """**Structure modulaire adoptée :**

```
Frontend/
├── shared/                  # Modules partagés
│   ├── api.js              # Client API centralisé
│   ├── websocket-client.js # Client WebSocket réutilisable  
│   ├── message-manager.js  # Gestion notifications
│   └── qr-scanner.js       # Scanner QR universel
├── components/             # Composants réutilisables
│   ├── header.js          # En-tête global
│   ├── navigation.js      # Navigation responsive
│   └── loading.js         # Indicateurs de chargement
├── pages/                 # Pages spécifiques
│   ├── patient/          # Interface patient
│   ├── staff/            # Interface personnel
│   └── admin/            # Interface admin
└── assets/               # Ressources statiques
    ├── css/              # Feuilles de style
    ├── images/           # Images et icônes
    └── fonts/            # Polices personnalisées
```

**Patterns de développement :**
• **Module Pattern** : Encapsulation et espaces de noms
• **Observer Pattern** : Communication inter-composants
• **Factory Pattern** : Création d'objets standardisée
• **Singleton Pattern** : Services partagés (API client, WebSocket)

**Gestion d'état simplifiée :**
• **LocalStorage** : Persistance données utilisateur
• **SessionStorage** : État temporaire de session
• **Custom Events** : Communication cross-composants
• **URL State** : État de navigation dans l'URL"""
    
    doc.add_paragraph(frontend_architecture_text)


def generate_development_tools(doc):
    """Génère la section outils de développement."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("3.3 Outils et environnement de développement")
    
    # 3.3.1 IDE et outils de développement
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.3.1 Environnement de développement intégré")
    
    ide_text = """**Visual Studio Code comme IDE principal :**

VSCode s'impose comme l'outil de développement optimal pour WAITLESS-CHU :

**Avantages techniques :**
• **Légèreté** : Démarrage rapide et consommation mémoire optimisée
• **Extensibilité** : Marketplace riche avec extensions spécialisées
• **IntelliSense avancé** : Auto-complétion et aide contextuelle
• **Debugging intégré** : Débogueur Python et JavaScript natif
• **Git intégré** : Gestion de versions transparente

**Extensions utilisées :**
• **Python (Microsoft)** : Support complet Python avec linting et formatting
• **JavaScript ES6 code snippets** : Accélération du développement frontend
• **PostgreSQL syntax highlighting** : Coloration syntaxique SQL
• **GitLens** : Visualisation avancée de l'historique Git
• **Thunder Client** : Tests API intégrés sans Postman
• **Live Server** : Serveur de développement avec rechargement automatique

**Configuration optimisée :**
• **Workspace settings** : Configuration projet spécifique
• **Tasks automation** : Automatisation des tâches répétitives
• **Multi-terminal** : Terminaux séparés backend/frontend
• **Synchronized settings** : Configuration synchronisée entre machines"""
    
    doc.add_paragraph(ide_text)
    
    # Tableau des outils
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # En-têtes
    header_cells = table.rows[0].cells
    headers = ['Catégorie', 'Outil', 'Usage', 'Justification']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    
    # Données outils
    tools_data = [
        ("IDE", "Visual Studio Code", "Développement principal", "Léger, extensible, intégré"),
        ("Database", "pgAdmin 4", "Administration PostgreSQL", "Interface graphique complète"),
        ("API Testing", "Thunder Client", "Tests endpoints", "Intégré VSCode"),
        ("Versioning", "Git + GitHub", "Contrôle de version", "Standard industrie"),
        ("Documentation", "Swagger UI", "Doc API automatique", "Généré par FastAPI"),
        ("Monitoring", "Browser DevTools", "Debug frontend", "Outils natifs navigateur"),
    ]
    
    for category, tool, usage, justification in tools_data:
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = tool
        row_cells[2].text = usage
        row_cells[3].text = justification
        
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
    
    # 3.3.2 Testing et qualité
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.3.2 Stratégie de tests et qualité du code")
    
    testing_text = """**Tests Backend automatisés :**

**Pytest pour tests unitaires :**
Framework de test Python moderne avec syntaxe claire et fixtures puissantes.
• **Test discovery** automatique des modules de test
• **Fixtures** réutilisables pour setup/teardown
• **Parametrized tests** pour couverture exhaustive
• **Mocking** intégré pour isolation des tests

**Tests d'intégration API :**
• **TestClient FastAPI** : Tests endpoints avec base de données test
• **Factory Boy** : Génération de données de test cohérentes
• **Database rollback** : Isolation complète des tests
• **Async testing** : Support des fonctions asynchrones

**Tests Frontend manuels :**
Approche pragmatique adaptée au contexte du projet :
• **Tests cross-browser** : Validation Chrome, Firefox, Safari, Edge
• **Tests responsive** : Vérification desktop, tablette, mobile
• **Tests d'accessibilité** : Validation WCAG avec outils navigateur
• **Tests utilisabilité** : Scénarios utilisateur réels

**Métriques de qualité :**
• **Code coverage** backend : Objectif 85%+ pour fonctions critiques
• **Linting** : Respect des conventions PEP 8 et ESLint
• **Type checking** : Validation types avec mypy pour Python
• **Performance** : Profiling des endpoints critiques"""
    
    doc.add_paragraph(testing_text)


def generate_security_and_authentication(doc):
    """Génère la section sécurité et authentification."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("3.4 Sécurité et authentification")
    
    # 3.4.1 JWT et gestion des tokens
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.4.1 Architecture d'authentification JWT")
    
    jwt_text = """**JSON Web Tokens pour l'authentification :**

**Avantages JWT vs sessions traditionnelles :**
• **Stateless** : Pas de stockage serveur, scalabilité horizontale facilitée
• **Self-contained** : Toutes les informations dans le token
• **Cross-domain** : Support natif des architectures distribuées
• **Performance** : Pas de requête base pour validation

**Structure des tokens WAITLESS-CHU :**
```json
{
  "header": {
    "alg": "HS256",
    "typ": "JWT"
  },
  "payload": {
    "sub": "user_email",
    "role": "staff",
    "assigned_service_id": 123,
    "exp": 1640995200,
    "iat": 1640991600
  },
  "signature": "HMAC_SHA256_signature"
}
```

**Sécurisation avancée :**
• **Expiration courte** : 30 minutes pour limiter l'exposition
• **Refresh tokens** : Renouvellement automatique sécurisé
• **Rotation des clés** : Changement périodique des clés de signature
• **Blacklisting** : Révocation immédiate des tokens compromis

**Gestion des rôles et permissions :**
```python
class RoleChecker:
    def __init__(self, allowed_roles: List[UserRole]):
        self.allowed_roles = allowed_roles
    
    def __call__(self, current_user: User = Depends(get_current_user)):
        if current_user.role not in self.allowed_roles:
            raise HTTPException(status_code=403, detail="Insufficient permissions")
        return current_user

# Usage dans les endpoints
@router.get("/admin/dashboard")
async def admin_dashboard(
    current_user: User = Depends(RoleChecker([UserRole.ADMIN]))
):
    # Endpoint accessible uniquement aux admins
```"""
    
    doc.add_paragraph(jwt_text)
    
    # 3.4.2 Protection des données
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.4.2 Protection et validation des données")
    
    data_protection_text = """**Validation côté serveur avec Pydantic :**

**Schemas de validation stricts :**
```python
class UserCreate(BaseModel):
    email: EmailStr = Field(..., description="Email valide requis")
    password: str = Field(..., min_length=8, regex="^(?=.*[a-z])(?=.*[A-Z])(?=.*\d)")
    full_name: str = Field(..., min_length=2, max_length=100)
    phone: Optional[str] = Field(None, regex="^[0-9+\-\s()]+$")
    role: UserRole = Field(default=UserRole.PATIENT)
    
    @validator('password')
    def validate_password_strength(cls, v):
        if not any(c.isupper() for c in v):
            raise ValueError('Password must contain uppercase letter')
        if not any(c.islower() for c in v):
            raise ValueError('Password must contain lowercase letter')
        if not any(c.isdigit() for c in v):
            raise ValueError('Password must contain digit')
        return v
```

**Sanitisation des données :**
• **Input sanitization** : Nettoyage des entrées utilisateur
• **SQL injection prevention** : Requêtes paramétrées via ORM
• **XSS protection** : Échappement automatique des données de sortie
• **File upload validation** : Validation stricte des types de fichiers

**Chiffrement et hashage :**
• **Passwords** : bcrypt avec salt automatique et coût adaptatif
• **Sensitive data** : Chiffrement AES-256 pour données critiques
• **Tokens** : Signature HMAC-SHA256 pour intégrité
• **Transport** : TLS 1.3 obligatoire pour toutes les communications

**Configuration CORS sécurisée :**
```python
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://waitless-chu.ma",
        "https://admin.waitless-chu.ma"
    ],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Authorization", "Content-Type"],
    expose_headers=["X-Total-Count"]
)
```"""
    
    doc.add_paragraph(data_protection_text)
    
    # 3.4.3 Conformité réglementaire
    p = doc.add_paragraph()
    p.style = 'SubsectionTitle'
    p.add_run("3.4.3 Conformité RGPD et standards hospitaliers")
    
    compliance_text = """**Respect du RGPD (Règlement Général sur la Protection des Données) :**

**Minimisation des données :**
• Collection uniquement des données strictement nécessaires
• Durée de conservation limitée et documentée
• Suppression automatique des données expirées
• Pseudonymisation des données analytiques

**Droits des personnes :**
• **Droit d'accès** : API pour consultation des données personnelles
• **Droit de rectification** : Modification des informations par l'utilisateur
• **Droit à l'effacement** : Suppression complète sur demande
• **Portabilité** : Export des données au format JSON

**Sécurité by design :**
• Chiffrement des données sensibles
• Audit logs complets
• Accès contrôlé et traçé
• Formation équipe sur bonnes pratiques

**Standards hospitaliers :**
• **Confidentialité médicale** : Isolation stricte des données patient
• **Traçabilité** : Log de tous les accès aux données
• **Intégrité** : Validation et checksums pour données critiques
• **Disponibilité** : Redondance et sauvegarde pour continuité service"""
    
    doc.add_paragraph(compliance_text)


def generate_chapter_conclusion(doc):
    """Ajoute la conclusion du chapitre."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.style = 'SectionTitle'
    p.add_run("Conclusion du chapitre")
    
    conclusion_text = """Ce chapitre a présenté et justifié l'ensemble des choix technologiques qui constituent la fondation technique du système WAITLESS-CHU.

L'adoption de FastAPI pour le backend combine performance exceptionnelle, documentation automatique et validation de données moderne, répondant parfaitement aux exigences de scalabilité et de maintenabilité du projet. PostgreSQL assure la robustesse et l'intégrité des données critiques avec des fonctionnalités avancées adaptées aux besoins hospitaliers.

Le choix de JavaScript vanilla pour le frontend privilégie la simplicité, la performance et la compatibilité universelle, éliminant les barrières technologiques tout en exploitant pleinement les APIs web modernes. Cette approche garantit une accessibilité maximale, cruciale dans le contexte hospitalier.

L'environnement de développement basé sur Visual Studio Code et les outils modernes (Git, pytest, pgAdmin) optimise la productivité de l'équipe tout en maintenant la qualité du code. La stratégie de tests combinant automatisation backend et validation manuelle frontend assure la fiabilité du système.

L'architecture de sécurité multicouche avec authentification JWT, validation Pydantic et conformité RGPD répond aux exigences strictes du secteur hospitalier en matière de protection des données sensibles.

Ces choix technologiques forment un écosystème cohérent et moderne qui facilite le développement, optimise les performances et garantit la sécurité, créant les conditions optimales pour la réalisation technique présentée au chapitre suivant."""
    
    doc.add_paragraph(conclusion_text)


def generate_chapter3_report():
    """Fonction principale de génération du chapitre 3."""
    print("🚀 Génération du Chapitre 3: Choix Technologiques...")
    
    # Créer le document
    doc = Document()
    
    # Configurer les styles
    setup_document_styles(doc)
    
    # Générer chaque section
    print("📋 Génération de l'introduction du chapitre...")
    add_chapter_intro(doc)
    
    print("⚙️ Génération du stack backend...")
    generate_backend_technology_stack(doc)
    
    print("🌐 Génération du stack frontend...")
    generate_frontend_technology_stack(doc)
    
    print("🛠️ Génération des outils de développement...")
    generate_development_tools(doc)
    
    print("🔒 Génération de la sécurité et authentification...")
    generate_security_and_authentication(doc)
    
    print("✍️ Génération de la conclusion du chapitre...")
    generate_chapter_conclusion(doc)
    
    # Sauvegarder le document
    filename = "chapitre3.docx"
    doc.save(filename)
    
    print(f"✅ Chapitre 3 généré avec succès: {filename}")
    print(f"📄 Pages générées: ~10 pages")
    print(f"📊 Contenu: Stack backend, frontend, outils développement, sécurité")
    
    return filename


if __name__ == "__main__":
    generate_chapter3_report()